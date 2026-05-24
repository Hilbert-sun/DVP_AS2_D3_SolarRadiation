from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "demo_script_revised_2-3min.docx"


def set_run_font(run, size=11, bold=False):
    run.font.name = "Arial"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run.font.size = Pt(size)
    run.bold = bold


def add_paragraph(doc, text="", size=11, bold=False, align=None, space_after=6):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = 1.12
    run = p.add_run(text)
    set_run_font(run, size=size, bold=bold)
    return p


def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8 if level == 1 else 5)
    p.paragraph_format.space_after = Pt(5)
    run = p.add_run(text)
    set_run_font(run, size=14 if level == 1 else 12, bold=True)
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(style=None)
    p.paragraph_format.left_indent = Cm(0.55)
    p.paragraph_format.first_line_indent = Cm(-0.25)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run("• " + text)
    set_run_font(run, size=10.5)
    return p


def add_rule(paragraph):
    p = paragraph._p
    p_pr = p.get_or_add_pPr()
    borders = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "999999")
    borders.append(bottom)
    p_pr.append(borders)


def build():
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Cm(1.7)
    section.bottom_margin = Cm(1.7)
    section.left_margin = Cm(1.8)
    section.right_margin = Cm(1.8)

    styles = doc.styles
    styles["Normal"].font.name = "Arial"
    styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    styles["Normal"].font.size = Pt(11)

    title = add_paragraph(
        doc,
        "FIT5147 DVP Part 1 Demo Speaking Script",
        size=17,
        bold=True,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        space_after=2,
    )
    add_rule(title)
    add_paragraph(
        doc,
        "When Sunshine Meets Rain: Understanding Solar Radiation Patterns at King's Park, Hong Kong",
        size=11,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        space_after=10,
    )

    add_heading(doc, "2-3分钟直接读稿版", 1)
    script = [
        "Good morning everyone. My project is called When Sunshine Meets Rain. It focuses on daily climate observations at King's Park in Hong Kong, especially global solar radiation, rainfall, sunshine duration, and relative humidity.",
        "The aim of this visualisation is to communicate how solar radiation changes over time, how it follows seasonal patterns, and how extreme high-radiation days are related to sunshine, rainfall, and humidity. My intended audience is Hong Kong urban planners and environmental policy advisors. They may not need detailed statistical modelling, but they do need clear visual evidence that can support solar planning and climate-adaptation decisions.",
        "In Sheet 1, I brainstormed possible ways to tell this story. I considered annual trend lines, monthly rhythm charts, GSR boxplots, SUN versus GSR scatterplots, seasonal facets, LOESS curves, season cards, and extreme-day summaries. I removed ideas such as a map because the data only comes from one station, and I removed raw daily lines and pairplots because they would be too noisy for this audience.",
        "Sheets 2, 3, and 4 explore three different complete designs. Sheet 2 is a scrollytelling design. It gives strong narrative guidance, but it is less flexible for comparison. Sheet 3 is a guided story dashboard with story steps, a main view, an insight panel, and filters. This gives a balance between explanation and exploration. Sheet 4 focuses more strongly on seasonal and extreme-event exploration, which is useful for planners but gives less emphasis to the long-term trend.",
        "For the final realisation in Sheet 5, I combine the strengths of Sheet 3 and Sheet 4. The interface uses story buttons to guide the reader through five stages: long-term context, seasonal rhythm, sunshine as the main driver, extreme solar radiation days, and planning takeaways. The main chart area shows the evidence, the insight panel explains the message, and the filters allow users to compare years, seasons, and extreme days.",
        "The key design idea is that this should not feel like a normal dashboard full of unrelated charts. It should work as a guided narrative. The user first understands the long-term and seasonal context, then sees the relationship between sunshine duration and solar radiation, and finally looks at extreme days and planning implications.",
        "The expected final implementation will use D3.js because it gives me more control over custom interactions, tooltips, story-step navigation, and linked highlighting between views. Overall, the design is intended to make the climate patterns readable, evidence-based, and useful for the planning audience. Thank you.",
    ]
    for paragraph in script:
        add_paragraph(doc, paragraph, size=11, space_after=7)

    add_heading(doc, "如果想控制在2分钟，可以删掉这两句", 2)
    add_bullet(doc, "The main chart area shows the evidence, the insight panel explains the message, and the filters allow users to compare years, seasons, and extreme days.")
    add_bullet(doc, "Overall, the design is intended to make the climate patterns readable, evidence-based, and useful for the planning audience.")

    doc.add_page_break()
    add_heading(doc, "老师可能会问的问题与回答", 1)

    qas = [
        (
            "1. Why did you choose urban planners and policy advisors as the audience?",
            "Because the topic is connected to solar planning and climate adaptation. This audience needs clear evidence for decision-making, so I designed a guided narrative with concise explanations instead of a highly technical statistical dashboard.",
        ),
        (
            "2. Why did you remove the map idea?",
            "The dataset is from King's Park, which is a single meteorological station. A map would suggest spatial comparison, but the data does not support that. So I removed the map to avoid misleading the audience.",
        ),
        (
            "3. Why did you choose the guided story dashboard as the final design?",
            "It balances guidance and exploration. Scrollytelling is strong for narrative but less flexible for comparison, while a pure dashboard may feel too open-ended. The guided dashboard lets users follow the story while still filtering by year, season, and extreme days.",
        ),
        (
            "4. Why use D3.js instead of R Shiny?",
            "The final design needs custom story-step navigation, SVG-level control, tooltip behaviour, and linked highlighting between views. D3.js gives finer control over these interactions, so it fits the final design better.",
        ),
        (
            "5. How do the interactions support the narrative rather than just add decoration?",
            "The interactions are tied to specific tasks. The year range helps inspect long-term context, the season filter supports seasonal comparison, and the extreme-day toggle helps users focus on high-radiation conditions while still understanding the wider pattern.",
        ),
    ]

    for q, a in qas:
        add_paragraph(doc, q, size=11, bold=True, space_after=3)
        add_paragraph(doc, "Answer: " + a, size=10.8, space_after=8)

    add_heading(doc, "现场答问小技巧", 1)
    tips = [
        "如果老师问为什么不用某个图，回答重点是：数据是否支持、受众是否需要、是否会增加认知负担。",
        "如果老师问 implementation feasibility，强调 D3 can support custom story steps, tooltips, filters, and linked highlighting.",
        "如果老师问 design evolution，按 Sheet 1 -> Sheets 2/3/4 -> Sheet 5 的逻辑说，不要只描述最后界面。",
        "如果忘词，抓住三句话：audience is planners; design is guided narrative; interaction supports comparison and evidence tracing.",
    ]
    for tip in tips:
        add_bullet(doc, tip)

    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build()
